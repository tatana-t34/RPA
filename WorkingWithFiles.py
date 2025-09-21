import pandas as pd
import csv
import os
from datetime import datetime
from docx import Document
from openpyxl import load_workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfReader, PdfWriter
import io

def create_csv_example(filename="data.csv"):
    if os.path.exists(filename):
        print(f"CSV уже существует: {filename}")
        return filename

    data = [
        ["Имя", "Количество", "Цена"],
        ["Товар А", 5, 100],
        ["Товар Б", 3, 250],
        ["Товар В", 2, 500],
        ["Товар Г", 4, 75]
    ]

    with open(filename, 'w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerows(data)

    print(f"Создан CSV файл: {filename}")
    return filename


def process_excel_report(csv_file, output_dir):
    df = pd.read_csv(csv_file)
    df["Стоимость"] = df["Количество"] * df["Цена"]

    excel_file = os.path.join(output_dir, "отчет.xlsx")
    df.to_excel(excel_file, sheet_name="Данные", index=False)

    wb = load_workbook(excel_file)
    ws = wb["Данные"]
    last_row = ws.max_row
    ws["E1"] = "Итого:"
    ws["E2"] = f"=SUM(D2:D{last_row})"
    wb.save(excel_file)

    total_sum = df["стоимость"].sum()
    print(f"Создан Excel отчет: {excel_file}")
    return excel_file, total_sum


def create_word_report(total_sum, output_dir):
    doc = Document()
    doc.add_heading("Смета", 0)
    doc.add_paragraph("Отчет сгенерирован роботом.")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Показатель"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Общая стоимость"
    table.cell(1, 1).text = str(total_sum)

    word_file = os.path.join(output_dir, "отчет.docx")
    doc.save(word_file)
    print(f"Создан Word отчет: {word_file}")
    return word_file


def fill_pdf_template(template_path, total_sum, output_dir):
    packet = io.BytesIO()
    pdfmetrics.registerFont(TTFont("Arial", "arial.ttf"))
    can = canvas.Canvas(packet, pagesize=letter)
    can.setFont("Arial", 12)

    can.drawString(210, 620, f"{datetime.now().strftime('%d.%m.%Y')}")
    can.drawString(220, 590, f"{total_sum} руб.")
    can.save()
    packet.seek(0)

    existing_pdf = PdfReader(open(template_path, "rb"))
    new_pdf = PdfReader(packet)
    output = PdfWriter()

    page = existing_pdf.pages[0]
    page.merge_page(new_pdf.pages[0])
    output.add_page(page)

    filled_pdf = os.path.join(output_dir, "отчет.pdf")
    with open(filled_pdf, "wb") as f:
        output.write(f)
    print(f"Создан PDF отчет: {filled_pdf}")
    return filled_pdf


def main():
    print("ОТЧЕТНЫЙ РОБОТ")

    output_dir = input("Введите путь для сохранения файлов (Enter для текущей директории): ").strip()
    if not output_dir:
        output_dir = os.getcwd()
    os.makedirs(output_dir, exist_ok=True)

    try:
        #CSV
        csv_file = create_csv_example(os.path.join(output_dir, "исходные_данные.csv"))

        #Excel
        excel_file, total_sum = process_excel_report(csv_file, output_dir)

        #Word
        word_file = create_word_report(total_sum, output_dir)

        #PDF (требуется готовый шаблон blank.pdf в той же папке)
        template_path = os.path.join(output_dir, "blank.pdf")
        if not os.path.exists(template_path):
            raise FileNotFoundError("Не найден шаблон blank.pdf! Поместите его в папку вывода.")
        pdf_file = fill_pdf_template(template_path, total_sum, output_dir)

        print("\nОТЧЕТЫ УСПЕШНО СОЗДАНЫ")
        print(f"CSV: {csv_file}")
        print(f"Excel: {excel_file}")
        print(f"Word: {word_file}")
        print(f"PDF: {pdf_file}")

    except Exception as e:
        print("Ошибка:", e)


if __name__ == "__main__":
    # pip install pandas openpyxl python-docx reportlab PyPDF2
    main()
